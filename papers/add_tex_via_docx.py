#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# coding: utf-8

import sys
import re
import docx
import os

# NOTE: before use, write 'pip install python-docx' in console

def umlauts(str):
	#str = str.replace("\xe4", '\\"a')
	#str = str.replace("\xf6", '\\"o')
	#str = str.replace("\xfc", '\\"u')
	#str = str.replace("\xc4", '\\"A')
	#str = str.replace("\xd6", '\\"O')
	#str = str.replace("\xdc", '\\"U')
	str = str.replace("&", '\\&')
	return str
    
def forceUTF8(input):
    try:
        input.decode('utf-8')
    except UnicodeError:
        convertedInput = ""
        for char in input:
            try:
                char.decode('utf-8')
                convertedInput += char
            except UnicodeError:
                hexChar = hex(ord(char))
                if hexChar in charReplacementMap:
                    replacement = charReplacementMap[hexChar]
                    convertedInput += replacement
                else:
                    convertedInput += char
        input = convertedInput.decode('utf-8')
    return input


charReplacementMap = {"0x96": "-", "0xe9": "é", "0xc4": "Ä", "0xd6": "Ö", "0xdc": "Ü", "0xdf": "ß", "0xe4": "ä", "0xf6": "ö", "0xfc": "ü", "0xf1": "ñ"}
temp_tex = open("paper.tex.template").read()
sys.stdout = os.fdopen(sys.stdout.fileno(), 'w')
tex = u""
docPath = os.path.join(sys.argv[1], sys.argv[2])
print(docPath)
doc = docx.Document(docPath)
old_style = doc.paragraphs[0].style.name
style_index = 0
title = ""
author = ""
for p in doc.paragraphs:
	if style_index < 2:
		if p.style.name != old_style:
			style_index += 1
			old_style = p.style.name
		if style_index == 0:
			title += " " + p.text
		elif style_index == 1:
			author += " " + p.text
	else:
		break
author = author.strip()
title = title.strip()
title = title.replace('\r', '')
title = title.replace('\n', ' ')
title = re.sub('\s\s+', ' ', title)

author = re.sub(r"\s\s+", " ", author)
author = re.sub(r"\s?[,;]\s?", ", ", author)

for currLine in temp_tex.split("\n"):
	m_author = re.match(".*\\\\author.*", currLine)
	m_title = re.match(".*\\\\title.*", currLine)
	if m_title:
		tex += "\\title{" + title + "}\n"
	elif m_author:
		tex += "\\author{" +  author + "}\n"
	else:
		tex += currLine + "\n"

open(sys.argv[1] + "/paper.tex",'w').write(tex)
